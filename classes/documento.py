import os.path
from os import path

from docx import Document
from docx.shared import Inches
from docx.shared import Mm
from docx.shared import Pt
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_BREAK

class Documento:
    document = None
    def __init__(self, src_template):
        if path.exists(src_template):
            self.document = Document(src_template)
        else:
            self.document = Document()
        style = self.document.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(9)

    def make_rows_bold(self, cells):
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    def cellStyle(self, cells, background):
        self.make_rows_bold(cells)
        for cell in cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcVAlign = OxmlElement("w:shd")
            tcVAlign.set(qn("w:fill"), background)
            tcPr.append(tcVAlign)

    def buildTableSection(self, doc_table, json):
        for row in json:
            tr_json = json[row]
            try:
                tr = doc_table.rows[row].cells
            except:
                tr = doc_table.add_row().cells
            if tr_json["style"] == "bold":
                self.cellStyle(tr, "#D9D9D9")
            for t in tr_json["text"]:
                tr[t].text = tr_json["text"][t]
    def mergeCells(self, doc_table, json):
        for merge in json:
            a = doc_table.cell(merge["a"][0], merge["a"][1])
            b = doc_table.cell(merge["b"][0], merge["b"][1])
            a.merge(b)
    def createTable(self, json):
        doc_table = self.document.add_table(json["rows"], json["cols"])
        self.buildTableSection(doc_table, json["thead"])
        self.buildTableSection(doc_table, json["tbody"])
        if "merge" in json:
            self.mergeCells(doc_table, json["merge"])
    def createHeading(self, text):
        self.document.add_heading(text, level=1)

    def createParagraph(self, text):
        self.document.add_paragraph(text)

    def save(self, output):
        self.document.save(output)

    def bytesToMb(self, bytes, force=False, level=0):
        if (force):
            force
        else:
            while bytes > 1024:
                bytes = bytes / 1024
                level += 1
        if str(bytes).find(".") > 0:
            bytes = "{:.2f}".format(bytes)
        elif str(bytes) == "None":
            bytes = 0
        return str(bytes) + " " + self.getBytesMesure(level)

    def getBytesMesure(self, level):
        if level == 0:
            return "Bytes"
        elif level == 1:
            return "KB"
        elif level == 2:
            return "MB"
        elif level == 3:
            return "GB"
        return ""